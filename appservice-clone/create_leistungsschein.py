from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def add_bottom_border(paragraph):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFBFBF')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_field_code(run, field_name):
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_name

    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')

    text = OxmlElement('w:t')
    text.text = '1'

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def sanitize_filename_part(value, fallback):
    if not value:
        return fallback
    cleaned = re.sub(r'[<>]', '', value).strip()
    cleaned = re.sub(r'\s+', '_', cleaned)
    cleaned = re.sub(r'[^A-Za-z0-9_-]', '_', cleaned)
    cleaned = re.sub(r'_+', '_', cleaned).strip('_')
    return cleaned or fallback


def apply_run_font(paragraph, size=11, bold=False):
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.bold = bold


def main():
    today_str = date.today().isoformat()

    data = {
        'kunde': '<<Kundenname>>',
        'ansprechpartner_kunde': '<<Name, Rolle, E-Mail>>',
        'it_dienstleister': '<<Firmenname>>',
        'ansprechpartner_dienstleister': '<<Name, Rolle, E-Mail>>',
        'projektname': '<<Projektname>>',
        'version': '1.0',
        'datum': today_str,
        'gueltigkeitszeitraum': 'bis auf Widerruf',
        'leistungsort': 'Hybrid, <<Adresse>>',
        'abrechnungsmodell': 'T&M',
        'tagessatz_stundensatz': '<<EUR>>',
        'sla': 'Mo-Fr 09-17 Uhr',
        'abhaengigkeiten': 'Tenant-Zugaenge, VPN, Admin-Rechte',
    }

    placeholders = set()

    def track(text):
        if isinstance(text, str):
            placeholders.update(re.findall(r'<<[^<>]+>>', text))

    def add_heading(text, level=1):
        track(text)
        p = doc.add_heading(text, level=level)
        apply_run_font(p, size=17 if level == 1 else 13)
        return p

    def add_text(text, style=None):
        track(text)
        p = doc.add_paragraph(text, style=style)
        apply_run_font(p, size=11)
        return p

    doc = Document()

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    for style_name, size in [('Heading 1', 17), ('Heading 2', 13), ('Heading 3', 12)]:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('Leistungsschein')
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)

    project_line = doc.add_paragraph()
    project_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    project_text = f"Projekt: {data['projektname']} | Kunde: {data['kunde']} | Version: {data['version']} | Datum: {data['datum']}"
    track(project_text)
    project_run = project_line.add_run(project_text)
    project_run.font.name = 'Calibri'
    project_run.font.size = Pt(11)

    separator = doc.add_paragraph()
    add_bottom_border(separator)

    add_heading('Dokumentmetadaten', level=1)

    metadata_rows = [
        ('Kunde', data['kunde']),
        ('Ansprechpartner Kunde', data['ansprechpartner_kunde']),
        ('IT-Dienstleister', data['it_dienstleister']),
        ('Ansprechpartner Dienstleister', data['ansprechpartner_dienstleister']),
        ('Projektname', data['projektname']),
        ('Version / Datum', f"{data['version']} / {data['datum']}"),
        ('Gueltigkeitszeitraum', data['gueltigkeitszeitraum']),
        ('Leistungsort', data['leistungsort']),
        ('Abrechnungsmodell', data['abrechnungsmodell']),
    ]

    metadata_table = doc.add_table(rows=0, cols=2)
    metadata_table.style = 'Table Grid'
    for label, value in metadata_rows:
        track(label)
        track(value)
        row_cells = metadata_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    add_heading('Einleitung', level=1)
    add_text(
        'Dieser Leistungsschein beschreibt Ziel, Inhalt und Rahmenbedingungen der beauftragten IT-Leistungen '
        'im Projektkontext und dient als verbindliche Grundlage fuer die operative Umsetzung.'
    )
    add_text(
        'Er grenzt den vereinbarten Leistungsumfang gegen nicht enthaltene Leistungen ab und wird zur '
        'Steuerung von Durchfuehrung, Nachverfolgung und Abnahme der Leistung verwendet.'
    )

    add_heading('Leistungsumfang / Leistungen', level=1)
    add_text('Die nachfolgenden Leistungen sind Bestandteil dieses Leistungsscheins:')

    leistungen = [
        'Ist-Analyse, Kickoff und Anforderungsaufnahme der aktuellen Umgebung.',
        'Konzeption einer Zielarchitektur beziehungsweise eines Zielbilds fuer <<Plattform/Workload>>.',
        'Einrichtung und Mandantenkonfiguration (z. B. M365, Azure, Netzwerk) gemaess abgestimmtem Scope.',
        'Umsetzung einer Security-Baseline inklusive Hardening und Richtlinien.',
        'Implementierung, Migration und Rollout fuer <<Systeme/Benutzergruppen>>.',
        'Technische Dokumentation (Betriebshandbuch, Architektur, Uebergabeunterlagen).',
        'Wissenstransfer und Uebergabe-Workshop mit den Betriebsverantwortlichen.',
        'Durchfuehrung von Tests und Unterstuetzung bei der Abnahme.',
        'Betriebsunterstuetzung (Hypercare) fuer <<X Tage>> nach Go-Live.',
    ]
    for item in leistungen:
        add_text(item, style='List Bullet')

    add_heading('Liefergegenstaende', level=2)
    liefergegenstaende = [
        'Aktualisierte technische Dokumentation.',
        'Konfigurations- und Umsetzungsnachweise.',
        'Uebergabeprotokoll inkl. offener Punkte.',
    ]
    for item in liefergegenstaende:
        add_text(item, style='List Bullet')

    add_heading('Nicht im Leistungsumfang', level=1)
    nicht_enthalten = [
        'Enduser-Support ausserhalb des definierten Scopes.',
        'Lizenzkosten und laufende Subskriptionsgebuehren von Herstellern.',
        'Beschaffung und Austausch von Hardware.',
        'Leistungen von Drittanbietern ohne separate Beauftragung.',
        'Schulungen ausserhalb des vereinbarten Wissenstransfers.',
    ]
    for item in nicht_enthalten:
        add_text(item, style='List Bullet')

    add_heading('Annahmen', level=1)
    annahmen = [
        'Der Kunde stellt benoetigte Zugaenge, Rechte und Systeme rechtzeitig bereit.',
        'Benannte Ansprechpartner stehen fuer Termine, Reviews und Entscheidungen verfuegbar zur Verfuegung.',
        'Die bestehende Umgebung entspricht dem dokumentierten Ist-Stand.',
        'Entscheidungen und Abnahmen erfolgen innerhalb von <<X Werktagen>> nach Vorlage.',
        f"Abhaengigkeiten werden aktiv adressiert: {data['abhaengigkeiten']}.",
    ]
    for item in annahmen:
        add_text(item, style='List Bullet')

    add_heading('Einschraenkungen', level=1)
    einschraenkungen = [
        'Aenderungen am Scope erfolgen ausschliesslich ueber einen Change-Request-Prozess.',
        'Leistungserbringung kann von Drittanbieter-Verfuegbarkeiten und Provider-SLAs abhaengen.',
        'Zeitplaene sind von Freigaben, Zugaengen und Mitwirkungen des Kunden abhaengig.',
        'Fuer Legacy-Systeme ausserhalb Herstellersupport besteht keine Gewaehr fuer Fehlerbehebung.',
    ]
    for item in einschraenkungen:
        add_text(item, style='List Bullet')

    add_heading('Mitwirkungspflichten Kunde', level=1)
    mitwirkung = [
        'Bereitstellung erforderlicher Daten, Informationen und Systemzugaenge.',
        'Teilnahme an Abstimmungen, Workshops und Review-Terminen.',
        'Fristgerechte Durchfuehrung von Tests und Rueckmeldung zu Ergebnissen.',
        'Benennung eines entscheidungsbefugten Hauptansprechpartners.',
    ]
    for item in mitwirkung:
        add_text(item, style='List Bullet')

    add_heading('Termine & Meilensteine', level=1)
    milestone_table = doc.add_table(rows=1, cols=3)
    milestone_table.style = 'Table Grid'
    hdr_cells = milestone_table.rows[0].cells
    hdr_cells[0].text = 'Meilenstein'
    hdr_cells[1].text = 'Beschreibung'
    hdr_cells[2].text = 'Zieltermin'

    milestones = [
        ('Kickoff', 'Projektstart und Scope-Abstimmung', '<<YYYY-MM-DD>>'),
        ('Technische Umsetzung', 'Konfiguration und Implementierung im vereinbarten Umfang', '<<YYYY-MM-DD>>'),
        ('Abnahme', 'Abschlusstest, Uebergabe und Abnahmeprotokoll', '<<YYYY-MM-DD>>'),
    ]
    for milestone, description, due_date in milestones:
        track(milestone)
        track(description)
        track(due_date)
        row = milestone_table.add_row().cells
        row[0].text = milestone
        row[1].text = description
        row[2].text = due_date

    add_heading('Abnahme', level=1)
    add_text(
        'Die Abnahme erfolgt anhand der dokumentierten Abnahmekriterien. Der Auftraggeber prueft die '
        'Leistungsergebnisse innerhalb einer Frist von <<X Werktagen>> und dokumentiert das Ergebnis im '
        'Abnahmeprotokoll.'
    )

    add_heading('Preise & Zahlungsbedingungen', level=1)
    model = data['abrechnungsmodell'].strip().lower()
    if model == 'festpreis':
        preistext = 'Festpreis: <<EUR>> netto. Zahlungsplan: 50 % bei Beauftragung, 50 % nach Abnahme.'
    elif model == 'paketpreis':
        preistext = 'Paketpreis: <<EUR>> netto gemaess vereinbartem Leistungspaket.'
    else:
        preistext = (
            f"T&M-Modell: Verrechnung nach Aufwand zum Satz von {data['tagessatz_stundensatz']} je Stunde/Tag, "
            'Abrechnung monatlich nach Leistungsnachweis.'
        )
    add_text(preistext)
    add_text('Reise- und Nebenkosten werden gemaess <<Regelung>> abgerechnet.')

    add_heading('Vertraulichkeit & Datenschutz', level=1)
    add_text(
        'Beide Parteien behandeln alle im Rahmen dieses Leistungsscheins ausgetauschten Informationen vertraulich '
        'und verwenden sie ausschliesslich zur Vertragserfuellung.'
    )
    add_text(
        'Personenbezogene Daten werden nur im erforderlichen Umfang und auf Basis der geltenden Datenschutzvorgaben '
        'verarbeitet. Sofern vorhanden, gelten zusaetzlich bestehende NDA- und Auftragsverarbeitungsvereinbarungen.'
    )
    add_text('Dieser Leistungsschein enthaelt keine Rechtsberatung.')

    add_heading('Unterschriften', level=1)
    sign_table = doc.add_table(rows=4, cols=2)
    sign_table.style = 'Table Grid'
    sign_table.cell(0, 0).text = 'Fuer Kunde'
    sign_table.cell(0, 1).text = 'Fuer IT-Dienstleister'
    sign_table.cell(1, 0).text = f"Name: {data['ansprechpartner_kunde']}"
    sign_table.cell(1, 1).text = f"Name: {data['ansprechpartner_dienstleister']}"
    sign_table.cell(2, 0).text = 'Datum: <<YYYY-MM-DD>>'
    sign_table.cell(2, 1).text = 'Datum: <<YYYY-MM-DD>>'
    sign_table.cell(3, 0).text = 'Unterschrift: ____________________'
    sign_table.cell(3, 1).text = 'Unterschrift: ____________________'

    for row in sign_table.rows:
        for cell in row.cells:
            track(cell.text)

    section = doc.sections[0]
    header = section.header
    if not header.paragraphs:
        header_p = header.add_paragraph()
    else:
        header_p = header.paragraphs[0]
    header_text = f"Leistungsschein - {data['projektname']}"
    track(header_text)
    header_p.text = header_text
    header_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in header_p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(9)

    footer = section.footer
    if not footer.paragraphs:
        footer_p = footer.add_paragraph()
    else:
        footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_prefix = footer_p.add_run('Seite ')
    run_prefix.font.name = 'Calibri'
    run_prefix.font.size = Pt(9)

    run_page = footer_p.add_run()
    run_page.font.name = 'Calibri'
    run_page.font.size = Pt(9)
    add_field_code(run_page, 'PAGE')

    run_mid = footer_p.add_run(' von ')
    run_mid.font.name = 'Calibri'
    run_mid.font.size = Pt(9)

    run_total = footer_p.add_run()
    run_total.font.name = 'Calibri'
    run_total.font.size = Pt(9)
    add_field_code(run_total, 'NUMPAGES')

    run_suffix = footer_p.add_run(f" | Version {data['version']}")
    run_suffix.font.name = 'Calibri'
    run_suffix.font.size = Pt(9)

    kunde_part = sanitize_filename_part(data['kunde'], 'Kunde')
    projekt_part = sanitize_filename_part(data['projektname'], 'Projekt')
    date_part = data['datum'] if re.fullmatch(r'\d{4}-\d{2}-\d{2}', data['datum']) else today_str

    filename = f"Leistungsschein_{kunde_part}_{projekt_part}_{date_part}.docx"
    output_path = Path.cwd() / filename
    doc.save(output_path)

    print(f"DOCX erstellt: {output_path}")
    print('Review-Liste (offene Platzhalter):')
    if placeholders:
        for ph in sorted(placeholders):
            print(f"- {ph}")
    else:
        print('- Keine offenen Platzhalter gefunden.')


if __name__ == '__main__':
    main()
