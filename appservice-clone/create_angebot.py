from __future__ import annotations

from datetime import date
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


PH_RE = re.compile(r"<<[^<>]+>>")


def sanitize_file_part(value: str, fallback: str) -> str:
    text = re.sub(r"[<>]", "", value or "").strip()
    text = re.sub(r"\s+", "_", text)
    text = re.sub(r"[^A-Za-z0-9_-]", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text or fallback


def collect_placeholders(value: Any, target: set[str]) -> None:
    if isinstance(value, str):
        target.update(PH_RE.findall(value))
    elif isinstance(value, dict):
        for v in value.values():
            collect_placeholders(v, target)
    elif isinstance(value, list):
        for v in value:
            collect_placeholders(v, target)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(17)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)


def style_paragraph(paragraph, size: int = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold


def set_cell(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        style_paragraph(p, 11)


def add_field(run, field_name: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, txt, end])


def add_header_footer(doc: Document, project: str, version: str) -> None:
    for section in doc.sections:
        hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        hp.text = f"Angebot - {project}"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(hp, 9)

        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("Seite ")
        r2 = fp.add_run()
        r3 = fp.add_run(" von ")
        r4 = fp.add_run()
        r5 = fp.add_run(f" | Version {version}")
        add_field(r2, "PAGE")
        add_field(r4, "NUMPAGES")
        for r in [r1, r2, r3, r4, r5]:
            r.font.name = "Calibri"
            r.font.size = Pt(9)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    style_paragraph(p, 17 if level == 1 else 13, bold=True)


def text(doc: Document, value: str, style: str | None = None) -> None:
    p = doc.add_paragraph(value, style=style)
    style_paragraph(p, 11)


def bullets(doc: Document, values: list[str]) -> None:
    for item in values:
        text(doc, item, style="List Bullet")


def build_ls(doc: Document, meta: dict[str, str], ls: dict[str, Any], with_title: bool) -> None:
    if with_title:
        p = doc.add_paragraph(f"{ls['id']} - {ls['titel']}")
        style_paragraph(p, 17, True)

    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"
    rows = [
        ("LS-ID", ls["id"]),
        ("Titel", ls["titel"]),
        ("Pflicht/Optional", ls["typ"]),
        ("Aufwand (PT)", str(ls["pt"])),
        ("Preis", ls["preis"]),
        ("Version / Datum", f"{meta['version']} / {meta['datum']}"),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell(t.cell(i, 0), k)
        set_cell(t.cell(i, 1), v)

    heading(doc, "Einleitung", 2)
    text(doc, ls["einleitung"])
    heading(doc, "Leistungen", 2)
    bullets(doc, ls["leistungen"])
    heading(doc, "Annahmen", 2)
    bullets(doc, ls["annahmen"])
    heading(doc, "Einschraenkungen", 2)
    bullets(doc, ls["einschraenkungen"])
    heading(doc, "Liefergegenstaende", 2)
    bullets(doc, ls["liefergegenstaende"])
    heading(doc, "Nicht im Leistungsumfang", 2)
    bullets(doc, ls["out_of_scope"])
    heading(doc, "Aufwand & Abrechnung", 2)
    text(doc, ls["aufwand_abrechnung"])
    heading(doc, "Abnahmekriterien", 2)
    text(doc, ls["abnahme"])


def build_offer(
    doc: Document,
    meta: dict[str, str],
    scenario: dict[str, str],
    phases: list[dict[str, str]],
    milestones: list[dict[str, str]],
    risks: list[dict[str, str]],
    ls_list: list[dict[str, Any]],
    ls_files: dict[str, str],
) -> None:
    p = doc.add_paragraph("Angebot - Migration Netzwerklaufwerke nach Azure Storage")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, 18, True)

    meta_tab = doc.add_table(rows=5, cols=2)
    meta_tab.style = "Table Grid"
    title_rows = [
        ("Kunde", meta["kunde"]),
        ("Projektname", meta["projektname"]),
        ("Version", meta["version"]),
        ("Datum", meta["datum"]),
        ("Anbieter", meta["dienstleister"]),
    ]
    for i, (k, v) in enumerate(title_rows):
        set_cell(meta_tab.cell(i, 0), k)
        set_cell(meta_tab.cell(i, 1), v)

    heading(doc, "2) Management Summary")
    text(
        doc,
        f"Der Windows Fileserver {scenario['source']} mit ca. {scenario['volume']} und {scenario['files']} Dateien "
        f"fuer rund {scenario['users']} Nutzer wird nach Azure Files migriert.",
    )
    text(doc, f"Zielbild: Azure Files (SMB) + Private Endpoint + Backup + Monitoring, Zugriff ueber {scenario['network']}.")
    text(doc, "AD-basierte Berechtigungen werden uebernommen, NTFS-ACLs werden soweit moeglich gemappt.")
    text(doc, "Abweichungen werden transparent dokumentiert und als Entscheidungspunkt vorgelegt.")
    total_pt = sum(int(item["pt"]) for item in ls_list)
    text(doc, f"Gesamtaufwand: {total_pt} PT (Schaetzung). Gesamtpreis: <<Gesamtpreis in EUR>>.")

    heading(doc, "3) Ausgangslage & Zielsetzung")
    text(doc, "In-Scope:")
    bullets(
        doc,
        [
            "Migration der Dateidaten inkl. geplanter Cutover auf Azure Files.",
            "Einrichtung von Private Endpoint, Azure Backup und Azure Monitor/Log Analytics.",
            "Nutzer- und Stakeholder-Kommunikation inkl. Pilot und Hypercare.",
        ],
    )
    text(doc, "Out-of-Scope:")
    bullets(
        doc,
        [
            "Ablosung nicht dateibasierter Applikationen.",
            "Langfristiger Betrieb ausserhalb des Hypercare-Fensters.",
            "Grosses Re-Design historischer Berechtigungsmodelle ausserhalb Scope.",
        ],
    )
    text(doc, "Entscheidungspunkte:")
    bullets(
        doc,
        [
            "Auth-Methode fuer Azure Files.",
            "Netzwerk-Anbindung (VPN, ExpressRoute, S2S-VPN).",
            "Share-Struktur, Namenskonzept und Rollback-Kriterien.",
        ],
    )

    heading(doc, "4) Vorgehensmodell / Phasenuebersicht")
    pt = doc.add_table(rows=1, cols=4)
    pt.style = "Table Grid"
    for i, h in enumerate(["Phase", "Inhalt", "Ergebnis", "LS-Referenzen"]):
        set_cell(pt.cell(0, i), h)
    for phase in phases:
        row = pt.add_row().cells
        set_cell(row[0], phase["phase"])
        set_cell(row[1], phase["inhalt"])
        set_cell(row[2], phase["ergebnis"])
        set_cell(row[3], phase["ls"])

    heading(doc, "5) Leistungsuebersicht")
    lt = doc.add_table(rows=1, cols=6)
    lt.style = "Table Grid"
    for i, h in enumerate(["LS-ID", "Titel", "Pflicht/Optional", "Aufwand (PT)", "Preis", "Kurzbeschreibung"]):
        set_cell(lt.cell(0, i), h)
    for ls in ls_list:
        row = lt.add_row().cells
        set_cell(row[0], ls["id"])
        set_cell(row[1], ls["titel"])
        set_cell(row[2], ls["typ"])
        set_cell(row[3], str(ls["pt"]))
        set_cell(row[4], ls["preis"])
        set_cell(row[5], ls["kurz"])

    heading(doc, "6) Zeitplan & Meilensteine")
    mt = doc.add_table(rows=1, cols=3)
    mt.style = "Table Grid"
    for i, h in enumerate(["Meilenstein", "Zieltermin", "Abhaengigkeiten"]):
        set_cell(mt.cell(0, i), h)
    for m in milestones:
        row = mt.add_row().cells
        set_cell(row[0], m["name"])
        set_cell(row[1], m["termin"])
        set_cell(row[2], m["deps"])

    heading(doc, "7) Mitwirkungspflichten Kunde")
    bullets(
        doc,
        [
            "Zugaenge, Freigaben, Testnutzer und Entscheidungswege rechtzeitig bereitstellen.",
            "Nutzerkommunikation und Wartungsfenster intern freigeben.",
            "Rueckmeldungen zu Tests, Abnahmen und offenen Punkten fristgerecht geben.",
        ],
    )

    heading(doc, "8) Annahmen & Einschraenkungen")
    bullets(
        doc,
        [
            "On-Prem AD ist verfuegbar und fuer die Zielauthentifizierung nutzbar.",
            "Netzwerkverbindung nach Azure ist vor Cutover stabil getestet.",
            "Aenderungen ausserhalb Scope nur ueber Change Request.",
        ],
    )

    heading(doc, "9) Risiken & Massnahmen")
    rt = doc.add_table(rows=1, cols=3)
    rt.style = "Table Grid"
    for i, h in enumerate(["Risiko", "Auswirkung", "Massnahme"]):
        set_cell(rt.cell(0, i), h)
    for r in risks:
        row = rt.add_row().cells
        set_cell(row[0], r["risk"])
        set_cell(row[1], r["impact"])
        set_cell(row[2], r["mitigation"])

    heading(doc, "10) Preise & Zahlungsbedingungen")
    text(doc, f"Abrechnungsmodell: {meta['abrechnung']}")
    text(doc, f"Stunden-/Tagessatz: {meta['satz']}")
    text(doc, f"Reise-/Nebenkosten: {meta['reise']}")
    text(doc, "Abrechnungsturnus: monatlich nach Leistungsnachweis bzw. gemaess Festpreis-Meilensteinen.")

    heading(doc, "11) Abnahme & Uebergabe")
    bullets(
        doc,
        [
            "Abnahme je LS anhand der festgelegten Abnahmekriterien.",
            "Uebergabe von Betriebsdoku, Monitoring-Setup und Backup/Restore-Guide.",
            "Restpunkte werden im Uebergabeprotokoll mit Termin dokumentiert.",
        ],
    )

    heading(doc, "12) Vertraulichkeit / Datenschutz")
    text(doc, "Vertrauliche Informationen werden nur zur Leistungserbringung genutzt.")
    text(doc, "Datenschutz wird vertragsgemaess umgesetzt; dieser Abschnitt ist keine Rechtsberatung.")

    heading(doc, "13) Unterschriften")
    st = doc.add_table(rows=4, cols=2)
    st.style = "Table Grid"
    set_cell(st.cell(0, 0), "Fuer den Kunden")
    set_cell(st.cell(0, 1), "Fuer den IT-Dienstleister")
    set_cell(st.cell(1, 0), f"Name: {meta['kunde_kontakt']}")
    set_cell(st.cell(1, 1), f"Name: {meta['dienstleister_kontakt']}")
    set_cell(st.cell(2, 0), "Datum: <<YYYY-MM-DD>>")
    set_cell(st.cell(2, 1), "Datum: <<YYYY-MM-DD>>")
    set_cell(st.cell(3, 0), "Unterschrift: ____________________")
    set_cell(st.cell(3, 1), "Unterschrift: ____________________")

    heading(doc, "14) Anhaenge")
    text(doc, "Leistungsscheine:")
    for ls in ls_list:
        text(doc, f"{ls['id']} - {ls['titel']} ({ls_files[ls['id']]})", style="List Bullet")


def main() -> None:
    today = date.today().isoformat()
    meta = {
        "kunde": "<<Kundenname>>",
        "kunde_kontakt": "<<Name, Rolle, E-Mail>>",
        "dienstleister": "<<Firmenname>>",
        "dienstleister_kontakt": "<<Name, Rolle, E-Mail>>",
        "projektname": "<<Azure File Migration>>",
        "version": "1.0",
        "datum": "<<YYYY-MM-DD>>",
        "gueltig": "<<bis auf Widerruf / Datum>>",
        "abrechnung": "<<Festpreis / T&M>>",
        "satz": "<<EUR>>",
        "reise": "<<inkl./zzgl.>>",
    }

    scenario = {
        "source": "FS01",
        "volume": "2 TB",
        "files": "350.000",
        "users": "120",
        "network": "<<VPN>>",
    }

    common_assumptions = [
        "Kunde stellt benoetigte Zugaenge und Ansprechpartner bereit.",
        "Freigaben und Entscheidungen erfolgen innerhalb <<X Werktagen>>.",
        "Aenderungen ausserhalb Scope erfolgen ueber Change Request.",
    ]

    common_limits = [
        "Provider-Limits und Drittabhaengigkeiten koennen Zeitplaene beeinflussen.",
        "Nicht dokumentierte Sonderfaelle koennen Zusatzabstimmungen ausloesen.",
    ]

    ls_list: list[dict[str, Any]] = [
        {
            "id": "LS-01",
            "titel": "Projektmanagement & Kickoff",
            "typ": "Pflicht",
            "pt": 4,
            "preis": "<<Preis LS-01 in EUR>>",
            "kurz": "Projektsetup, Governance und Stakeholder-Kommunikation.",
            "einleitung": "Steuert Projekt, Termine, Entscheidungen und Kommunikation.",
            "leistungen": [
                "Kickoff mit Scope, Rollen und Vorgehen.",
                "Status-/Risikoreporting und Steuertermine.",
                "Kommunikationsplan fuer Stakeholder und Nutzerankuendigung.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Kickoff-Protokoll", "Projektplan", "Kommunikationsplan"],
            "out_of_scope": ["Dauerhaftes PMO", "Organisationsentwicklung ausserhalb Projekt"],
            "aufwand_abrechnung": "4 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Kickoff-Protokoll und Projektplan sind freigegeben.",
        },
        {
            "id": "LS-02",
            "titel": "IST-Analyse & Bestandsaufnahme",
            "typ": "Pflicht",
            "pt": 6,
            "preis": "<<Preis LS-02 in EUR>>",
            "kurz": "Analyse von Shares, NTFS/ACL, Datenvolumen und Stakeholder-Bedarf.",
            "einleitung": "Liefert die belastbare Basis fuer Design und Migration.",
            "leistungen": [
                "Inventur von Shares, Datenmenge und Dateianzahl.",
                "Analyse NTFS-ACLs und AD-Gruppen.",
                "Stakeholder-Interviews zu Kritikalitaet und Nutzungsprofil.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["IST-Analysebericht", "Priorisierung kritischer Shares", "Offene-Punkte-Liste"],
            "out_of_scope": ["Vollstaendige Datenbereinigung", "Fachanwendungsanalyse ausserhalb Filescope"],
            "aufwand_abrechnung": "6 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Analysebericht ist fachlich und technisch abgenommen.",
        },
        {
            "id": "LS-03",
            "titel": "Zielkonzept & Migrationsplanung",
            "typ": "Pflicht",
            "pt": 6,
            "preis": "<<Preis LS-03 in EUR>>",
            "kurz": "Ziel-Share-Struktur, Cutover- und Rollback-Plan.",
            "einleitung": "Definiert Zielbild und Umsetzungslogik der Migration.",
            "leistungen": [
                "Design Share-Struktur und Namenskonzept.",
                "Cutover-, Freeze- und Rollback-Plan erstellen.",
                "Kommunikationsfahrplan fuer Umstellungsfenster.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Zielkonzept", "Migrationsrunbook", "Rollback-Entscheidungspunkte"],
            "out_of_scope": ["Fachliches Re-Design aller Datenstrukturen", "IAM-Transformation ausserhalb Files"],
            "aufwand_abrechnung": "6 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Zielkonzept und Cutover-Plan sind freigegeben.",
        },
        {
            "id": "LS-04",
            "titel": "Azure Grundlage / Landing-Zone-Readiness",
            "typ": "Pflicht",
            "pt": 5,
            "preis": "<<Preis LS-04 in EUR>>",
            "kurz": "Readiness fuer Netzwerk, Identity und Governance.",
            "einleitung": "Sichert die technischen Voraussetzungen fuer Azure Files.",
            "leistungen": [
                "Netzwerk-/DNS-Pruefung fuer Private Endpoint.",
                "Identity- und RBAC-Grundkonfiguration im Scope.",
                "Readiness-Check mit Restpunkteliste.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Readiness-Checkliste", "Freigabedokument", "RBAC-Nachweise"],
            "out_of_scope": ["Kompletter Landing-Zone-Neuaufbau", "SIEM-Einfuehrung"],
            "aufwand_abrechnung": "5 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Readiness ist technisch bestaetigt und dokumentiert.",
        },
        {
            "id": "LS-05",
            "titel": "Azure Storage / Azure Files Implementierung",
            "typ": "Pflicht",
            "pt": 5,
            "preis": "<<Preis LS-05 in EUR>>",
            "kurz": "Storage, Shares, Private Endpoint, Backup und Monitoring.",
            "einleitung": "Implementiert die vereinbarte Azure-Files-Zielplattform.",
            "leistungen": [
                "Storage Accounts und Shares bereitstellen.",
                "Private Endpoint und DNS einrichten.",
                "Backup und Monitoring aktivieren.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Betriebsbereite Zielplattform", "Konfigurationsnachweis", "Testprotokoll"],
            "out_of_scope": ["Cross-Region-Architektur", "Vollautomatisierung aller Deployments"],
            "aufwand_abrechnung": "5 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Zielplattform ist erreichbar, gesichert und monitorbar.",
        },
        {
            "id": "LS-06",
            "titel": "Berechtigungen & Zugriff",
            "typ": "Pflicht",
            "pt": 5,
            "preis": "<<Preis LS-06 in EUR>>",
            "kurz": "NTFS/ACL-Mapping, Gruppen, Zugriffstests, Abweichungsdoku.",
            "einleitung": "Stellt Berechtigungen fuer den produktiven Zugriff sicher.",
            "leistungen": [
                "Mapping von NTFS-ACLs auf Zielfreigaben soweit moeglich.",
                "AD-Gruppen zuordnen und Berechtigungspruefungen durchfuehren.",
                "Abweichungen dokumentieren und Entscheidungspunkte vorlegen.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Berechtigungsmatrix", "Testprotokolle", "Abweichungsliste"],
            "out_of_scope": ["Vollstaendige IAM-Neuordnung", "Compliance-Audit ausserhalb Scope"],
            "aufwand_abrechnung": "5 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Berechtigungsmatrix freigegeben und Zugriffstests erfolgreich.",
        },
        {
            "id": "LS-07",
            "titel": "Datenmigration & Cutover",
            "typ": "Pflicht",
            "pt": 6,
            "preis": "<<Preis LS-07 in EUR>>",
            "kurz": "Pre-Seed, Delta-Sync, Freeze-Window, Go-Live und Rollback.",
            "einleitung": "Fuehrt die technische Migration bis zur Umschaltung aus.",
            "leistungen": [
                "Pre-Seed und Delta-Sync mit Migrationsprotokoll.",
                "Cutover im Wartungsfenster inkl. Rollback-Bereitschaft.",
                "Nutzerkommunikation waehrend Umstellung und Go-Live.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Migrationslogs", "Go-Live-Protokoll", "Rollback-Checkliste"],
            "out_of_scope": ["Unbegrenzte Nachmigration", "Datenbereinigung ausserhalb Scope"],
            "aufwand_abrechnung": "6 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Cutover abgeschlossen und produktiver Zugriff auf Azure Files moeglich.",
        },
        {
            "id": "LS-08",
            "titel": "Tests & Abnahmeunterstuetzung",
            "typ": "Pflicht",
            "pt": 3,
            "preis": "<<Preis LS-08 in EUR>>",
            "kurz": "Testfaelle, Pilotgruppe, Feedbackschleife und Freigabeprozess.",
            "einleitung": "Validiert die Migration mit Fokus auf Nutzer- und Betriebsfaehigkeit.",
            "leistungen": [
                "Testkatalog erstellen und durchfuehren.",
                "Pilotgruppe koordinieren und Feedback auswerten.",
                "Abnahmeprozess dokumentiert unterstuetzen.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Testkatalog", "Pilot-Feedbackbericht", "Abnahmeprotokoll"],
            "out_of_scope": ["24/7 Testbegleitung", "UAT fuer fachfremde Systeme"],
            "aufwand_abrechnung": "3 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Kritische Testfaelle bestanden, Pilotgruppe gibt Freigabe.",
        },
        {
            "id": "LS-09",
            "titel": "Dokumentation & Wissenstransfer",
            "typ": "Pflicht",
            "pt": 3,
            "preis": "<<Preis LS-09 in EUR>>",
            "kurz": "Betriebshandbuch, Backup/Restore-Guide, Kurzschulung und How-To.",
            "einleitung": "Sichert den geordneten Wissenstransfer in den Regelbetrieb.",
            "leistungen": [
                "Betriebsdokumentation und Admin-Guide erstellen.",
                "Monitoring und Backup/Restore-Prozesse dokumentieren.",
                "Kurzschulung: neues Netzlaufwerk verbinden, Known Issues, Supportweg.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Betriebshandbuch", "Runbook", "Schulungsunterlagen"],
            "out_of_scope": ["Dauerhafte Schulungsserie", "Mehrsprachige Dokumentation"],
            "aufwand_abrechnung": "3 PT, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Dokumente sind uebergeben, Wissenstransfer wurde durchgefuehrt.",
        },
        {
            "id": "LS-10",
            "titel": "Hypercare / Betreuung",
            "typ": "Optional",
            "pt": 3,
            "preis": "<<Preis LS-10 in EUR>>",
            "kurz": "10 Werktage Stabilisierung mit Ticketkanal und optionalen Daily Check-ins.",
            "einleitung": "Stabilisiert den Betrieb in der unmittelbaren Nach-Go-Live-Phase.",
            "leistungen": [
                "Hypercare fuer 10 Werktage nach Go-Live.",
                "Supportkanal mit Triage und Priorisierung.",
                "Optionale taegliche Check-ins und Abschlussreview.",
            ],
            "annahmen": common_assumptions,
            "einschraenkungen": common_limits,
            "liefergegenstaende": ["Hypercare-Log", "Ticketuebersicht", "Abschlussreview"],
            "out_of_scope": ["Dauerhafter Managed Service", "24/7 Rufbereitschaft"],
            "aufwand_abrechnung": "3 PT optional, Abrechnung gemaess <<Festpreis / T&M>>.",
            "abnahme": "Hypercare abgeschlossen, kritische Incidents bearbeitet.",
        },
    ]

    phases = [
        {"phase": "Initialisierung", "inhalt": "Kickoff und Governance", "ergebnis": "Projektstart", "ls": "LS-01"},
        {"phase": "Analyse", "inhalt": "Bestandsaufnahme", "ergebnis": "Analysebericht", "ls": "LS-02"},
        {"phase": "Konzeption", "inhalt": "Zielkonzept und Planung", "ergebnis": "Freigegebenes Design", "ls": "LS-03"},
        {"phase": "Readiness/Build", "inhalt": "LZ-Readiness und Implementierung", "ergebnis": "Zielplattform", "ls": "LS-04, LS-05"},
        {"phase": "Access", "inhalt": "ACL-Mapping und Zugriffstests", "ergebnis": "Freigegebene Rechte", "ls": "LS-06"},
        {"phase": "Migration", "inhalt": "Pre-Seed, Delta, Cutover", "ergebnis": "Go-Live", "ls": "LS-07"},
        {"phase": "Test/Abnahme", "inhalt": "Pilot und Freigaben", "ergebnis": "Abnahme", "ls": "LS-08"},
        {"phase": "Transition", "inhalt": "Doku, Schulung, Hypercare", "ergebnis": "Regelbetrieb", "ls": "LS-09, LS-10"},
    ]

    milestones = [
        {"name": "Kickoff", "termin": "<<YYYY-MM-DD>>", "deps": "Beauftragung, Verfuegbarkeit"},
        {"name": "Zielkonzept freigegeben", "termin": "<<YYYY-MM-DD>>", "deps": "Analyse abgeschlossen"},
        {"name": "Cutover", "termin": "<<YYYY-MM-DD>>", "deps": "Readiness, Tests, Wartungsfenster"},
        {"name": "Abnahme", "termin": "<<YYYY-MM-DD>>", "deps": "Pilotgruppe und Doku"},
    ]

    risks = [
        {
            "risk": "ACL-Sonderfaelle im Altbestand",
            "impact": "Fehlende Zugriffe nach Go-Live",
            "mitigation": "Pilottests und dokumentierte Abweichungsentscheidung",
        },
        {
            "risk": "Netzwerkengpass im Cutover",
            "impact": "Verzoegerung",
            "mitigation": "Pre-Seed/Delta, Connectivity-Tests, Rollback-Option",
        },
        {
            "risk": "Unzureichende Nutzerkommunikation",
            "impact": "Erhoehtes Ticketaufkommen",
            "mitigation": "Ankuendigung, How-To, Hypercare-Kanal",
        },
    ]

    placeholders: set[str] = set()
    for obj in [meta, scenario, ls_list, phases, milestones, risks]:
        collect_placeholders(obj, placeholders)
    placeholders.add("<<Gesamtpreis in EUR>>")

    root = Path.cwd()
    ls_dir = root / "leistungsscheine"
    ls_dir.mkdir(parents=True, exist_ok=True)

    kunde = sanitize_file_part(meta["kunde"], "Kunde")
    projekt = sanitize_file_part(meta["projektname"], "Projekt")
    main_path = root / f"Angebot_{kunde}_{projekt}_{today}.docx"
    combo_path = root / f"Angebot_inkl_Leistungsscheine_{kunde}_{projekt}_{today}.docx"

    ls_files: dict[str, str] = {}
    ls_paths: list[str] = []

    for ls in ls_list:
        name = f"{ls['id']}_{sanitize_file_part(ls['titel'], 'Leistungsschein')}_{today}.docx"
        ls_files[ls["id"]] = name
        out = ls_dir / name
        d = Document()
        style_doc(d)
        add_header_footer(d, meta["projektname"], meta["version"])
        build_ls(d, meta, ls, with_title=True)
        d.save(out)
        ls_paths.append(str(out))

    offer_doc = Document()
    style_doc(offer_doc)
    add_header_footer(offer_doc, meta["projektname"], meta["version"])
    build_offer(offer_doc, meta, scenario, phases, milestones, risks, ls_list, ls_files)
    offer_doc.save(main_path)

    combo = Document()
    style_doc(combo)
    add_header_footer(combo, meta["projektname"], meta["version"])
    build_offer(combo, meta, scenario, phases, milestones, risks, ls_list, ls_files)
    combo.add_page_break()
    heading(combo, "Anhang - Leistungsscheine (Volltext)")
    for i, ls in enumerate(ls_list):
        heading(combo, f"{ls['id']} - {ls['titel']}", 2)
        build_ls(combo, meta, ls, with_title=False)
        if i < len(ls_list) - 1:
            combo.add_page_break()
    combo.save(combo_path)

    print("Erzeugte Dateien:")
    print(f"- {main_path}")
    print(f"- {combo_path}")
    for p in ls_paths:
        print(f"- {p}")

    print("\nReview-Liste (offene Platzhalter):")
    for ph in sorted(placeholders):
        print(f"- {ph}")


if __name__ == "__main__":
    main()
